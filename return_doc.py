from docx import Document
from io import BytesIO

def create(summarized_list):
    doc=Document()
    doc.add_heading("Summarized Doc",0)

    for i in list(summarized_list):
        doc.add_paragraph(i)
    
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream